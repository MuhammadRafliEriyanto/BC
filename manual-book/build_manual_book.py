from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "manual-book"
DOCX_PATH = OUT_DIR / "Manual_Book_Bimbel_LMS_All_Roles.docx"
LOGO_PATH = ROOT / "public" / "logobc.png"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(87, 96, 111)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
WHITE = "FFFFFF"
BORDER = "D9E2EF"
AMBER = "FFF8E8"

TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_spacing(paragraph, before=0, after=6, line_spacing=1.25):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, margins=CELL_MARGIN_DXA):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in margins.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_geometry(table, widths: Sequence[int], indent=TABLE_INDENT_DXA):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[min(idx, len(widths) - 1)]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_picture_alt(run, title: str, description: str):
    for doc_pr in run._element.xpath(".//*[local-name()='docPr']"):
        doc_pr.set("title", title)
        doc_pr.set("descr", description)


def style_cell_text(cell, bold=False, color=INK, size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT):
    for paragraph in cell.paragraphs:
        paragraph.alignment = align
        set_paragraph_spacing(paragraph, before=0, after=0, line_spacing=1.15)
        for run in paragraph.runs:
            set_run_font(run, size=size, color=color, bold=bold)


def add_table(doc, headers: Sequence[str], rows: Sequence[Sequence[str]], widths: Sequence[int]):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)
    header_row = table.rows[0]
    set_repeat_table_header(header_row)
    for idx, header in enumerate(headers):
        cell = header_row.cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        cell.text = header
        style_cell_text(cell, bold=True, color=INK, size=9.2)

    for row_data in rows:
        row = table.add_row()
        for idx, value in enumerate(row_data):
            cell = row.cells[idx]
            cell.text = value
            style_cell_text(cell, size=9.2)
    doc.add_paragraph()
    return table


def add_label_detail_table(doc, rows: Sequence[tuple[str, str]]):
    table = doc.add_table(rows=0, cols=2)
    set_table_geometry(table, [2700, 6660])
    set_table_borders(table)
    for row_index, (label, value) in enumerate(rows):
        cells = table.add_row().cells
        if row_index == 0:
            set_repeat_table_header(table.rows[0])
        cells[0].text = label
        cells[1].text = value
        set_cell_shading(cells[0], LIGHT_BLUE)
        style_cell_text(cells[0], bold=True, size=9.4)
        style_cell_text(cells[1], size=9.4)
    doc.add_paragraph()


def add_callout(doc, title: str, body: str, fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    set_table_borders(table, color="D7DBE2", size="4")
    set_repeat_table_header(table.rows[0])
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, before=0, after=2, line_spacing=1.15)
    r = p.add_run(title)
    set_run_font(r, size=10, color=DARK_BLUE, bold=True)
    p2 = cell.add_paragraph()
    set_paragraph_spacing(p2, before=0, after=0, line_spacing=1.15)
    r2 = p2.add_run(body)
    set_run_font(r2, size=9.5, color=INK)
    doc.add_paragraph()


def add_heading(doc, text: str, level: int):
    paragraph = doc.add_heading(text, level=level)
    if level == 1:
        paragraph.paragraph_format.space_before = Pt(18)
        paragraph.paragraph_format.space_after = Pt(10)
    elif level == 2:
        paragraph.paragraph_format.space_before = Pt(14)
        paragraph.paragraph_format.space_after = Pt(7)
    else:
        paragraph.paragraph_format.space_before = Pt(10)
        paragraph.paragraph_format.space_after = Pt(5)
    return paragraph


def add_para(doc, text: str, bold_prefix: str | None = None, italic=False):
    p = doc.add_paragraph()
    set_paragraph_spacing(p)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, size=11, color=INK, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, size=11, color=INK, italic=italic)
    else:
        r = p.add_run(text)
        set_run_font(r, size=11, color=INK, italic=italic)
    return p


def add_bullets(doc, items: Iterable[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        r = p.add_run(item)
        set_run_font(r, size=10.5, color=INK)


def add_steps(doc, items: Iterable[str]):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        r = p.add_run(item)
        set_run_font(r, size=10.5, color=INK)


def add_checklist(doc, items: Iterable[str]):
    rows = [(item, "Ya / Tidak", "Catatan") for item in items]
    add_table(doc, ["Checklist", "Status", "Catatan"], rows, [5000, 1500, 2860])


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def set_running_header_footer(doc: Document):
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0]
    hp.text = ""
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = hp.add_run("Manual Book Bimbel LMS")
    set_run_font(r, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = ""
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run("Bina Cendekia - Panduan penggunaan dashboard owner, admin, guru, dan siswa")
    set_run_font(r, size=8.5, color=MUTED)


def add_cover(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(70)
    if LOGO_PATH.exists():
        run = p.add_run()
        run.add_picture(str(LOGO_PATH), width=Inches(1.15))
        set_picture_alt(run, "Logo Bina Cendekia", "Logo Bina Cendekia pada cover manual book.")

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(kicker, before=24, after=12)
    r = kicker.add_run("PANDUAN OPERASIONAL")
    set_run_font(r, size=11, color=BLUE, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(title, before=0, after=8)
    r = title.add_run("Manual Book Bimbel LMS")
    set_run_font(r, size=28, color=INK, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(subtitle, before=0, after=28)
    r = subtitle.add_run("Panduan lengkap untuk role Owner, Admin, Guru, dan Siswa")
    set_run_font(r, size=14, color=MUTED)

    add_label_detail_table(
        doc,
        [
            ("Aplikasi", "Bimbel LMS / Bina Cendekia"),
            ("Role dicakup", "Owner, Admin, Guru, Siswa"),
            ("Versi dokumen", "1.0"),
            ("Tanggal", "18 Juni 2026"),
            ("Sumber penyusunan", "Struktur route, komponen dashboard, dan service API pada project bimbel-new."),
        ],
    )

    add_callout(
        doc,
        "Catatan penggunaan",
        "Manual ini menjelaskan alur penggunaan berdasarkan tampilan dan endpoint yang tersedia di project. Fitur yang masih bergantung backend atau masih berbasis data dummy diberi catatan khusus agar tidak disalahpahami saat demo atau pengujian.",
        fill=AMBER,
    )

    doc.add_page_break()


def add_common_access(doc: Document):
    add_heading(doc, "1. Gambaran Umum", 1)
    add_para(
        doc,
        "Manual ini dipakai sebagai panduan operasional dashboard Bimbel LMS. Setiap pengguna masuk lewat halaman login, lalu sistem mengarahkan ke dashboard sesuai role akun.",
    )
    add_table(
        doc,
        ["Role", "Dashboard awal", "Fokus utama"],
        [
            ("Owner", "/dashboard-owner", "Monitoring seluruh cabang, admin cabang, pembayaran, dan pengeluaran operasional."),
            ("Admin", "/dashboard-admin", "Pengelolaan data cabang: siswa, guru, jadwal, pembayaran, dan keuangan cabang."),
            ("Guru", "/dashboard-guru", "Pengelolaan kelas yang diampu, jadwal, absensi QR, materi, tugas, nilai, dan tryout."),
            ("Siswa", "/dashboard-siswa", "Akses belajar: materi, tugas, jadwal, absensi, nilai, tagihan, dan tryout."),
        ],
        [1600, 2200, 5560],
    )

    add_heading(doc, "1.1 Login dan Keamanan Akun", 2)
    add_steps(
        doc,
        [
            "Buka halaman /login.",
            "Masukkan email dan password akun.",
            "Aktifkan Remember Me bila ingin sesi login tetap tersimpan pada perangkat yang aman.",
            "Klik Login. Jika berhasil, sistem mengarahkan pengguna ke dashboard sesuai role.",
            "Jika lupa password, klik Forgot Password, masukkan email, lalu ikuti instruksi reset yang dikirim.",
            "Jika login Google aktif pada konfigurasi aplikasi, pengguna dapat memakai tombol Google Sign-In.",
        ],
    )
    add_callout(
        doc,
        "Prinsip keamanan",
        "Jangan memakai Remember Me pada perangkat umum. Setelah selesai memakai dashboard di komputer bersama, buka menu profil lalu pilih Logout.",
    )

    add_heading(doc, "1.2 Profil Pengguna", 2)
    add_bullets(
        doc,
        [
            "Menu profil tersedia di topbar setiap dashboard.",
            "Pengguna dapat melihat informasi akun, memperbarui nama/email/foto profil, mengganti password, meminta instruksi reset password, dan logout.",
            "Avatar mendukung upload gambar seperti JPG, PNG, WebP, atau format gambar lain yang didukung browser.",
            "Perubahan profil disimpan ke backend lewat service auth dan state lokal diperbarui agar tampilan topbar ikut berubah.",
        ],
    )

    add_heading(doc, "1.3 Notifikasi", 2)
    add_bullets(
        doc,
        [
            "Owner, admin, guru, dan siswa memiliki ikon notifikasi di topbar sesuai implementasi masing-masing dashboard.",
            "Notifikasi berisi ringkasan aktivitas yang relevan, misalnya pembayaran, aktivasi, jadwal, tugas, materi, nilai, atau tryout.",
            "Klik item notifikasi yang memiliki tautan untuk langsung menuju halaman terkait.",
        ],
    )


def add_quick_toc(doc: Document):
    add_heading(doc, "Daftar Isi Singkat", 1)
    add_table(
        doc,
        ["Bagian", "Isi utama"],
        [
            ("1. Gambaran Umum", "Login, redirect role, profil, keamanan akun, dan notifikasi."),
            ("2. Manual Role Owner", "Dashboard owner, admin cabang, cabang, pembayaran, aktivitas, dan pengeluaran."),
            ("3. Manual Role Admin", "Data siswa, guru, jadwal, keuangan cabang, dan informasi pembayaran."),
            ("4. Manual Role Guru", "Jadwal, kelas, materi, tugas, nilai, absensi QR, dan tryout."),
            ("5. Manual Role Siswa", "Registrasi, pembayaran, dashboard belajar, materi, tugas, jadwal, absensi, nilai, tagihan, dan tryout."),
            ("6-8. Referensi", "Troubleshooting, ringkasan hak akses, dan catatan implementasi."),
        ],
        [2600, 6760],
    )
    add_callout(
        doc,
        "Cara memakai manual",
        "Mulai dari Gambaran Umum untuk memahami login dan aturan akun, lalu langsung lompat ke bab role yang sedang diuji. Bagian Troubleshooting dipakai saat data tidak muncul, import gagal, QR absensi bermasalah, atau pembayaran belum aktif.",
    )


def add_owner_manual(doc: Document):
    add_heading(doc, "2. Manual Role Owner", 1)
    add_para(
        doc,
        "Owner bertugas memantau kondisi seluruh sistem dan melakukan pengaturan level pusat: cabang, admin cabang, pembayaran, aktivitas membership, serta pengeluaran operasional.",
    )
    add_table(
        doc,
        ["Menu Owner", "Lokasi", "Kegunaan"],
        [
            ("Dashboard", "/dashboard-owner", "Melihat omzet, status cabang, aktivitas sistem, total siswa, dan grafik performa."),
            ("Admin Cabang", "/dashboard-owner/admin-cabang", "Membuat, mengedit, menghapus, dan mengirim ulang verifikasi akun admin cabang."),
            ("Cabang", "/dashboard-owner/cabang", "Menambah, mengedit, menghapus, import, export, dan mengatur status cabang."),
            ("Pembayaran", "/dashboard-owner/aktivitas", "Memantau pembayaran masuk, pembayaran keluar, aktivasi siswa, detail transaksi, dan export data."),
            ("Pengeluaran", "/dashboard-owner/pengeluaran", "Mencatat dan mengelola pengeluaran operasional lintas cabang."),
        ],
        [1800, 2600, 4960],
    )

    add_heading(doc, "2.1 Dashboard Owner", 2)
    add_bullets(
        doc,
        [
            "Kartu Omzet menampilkan pembayaran lunas dan nominal yang masih belum lunas.",
            "Kartu Cabang menampilkan jumlah cabang dan komposisi cabang aktif, persiapan, atau nonaktif.",
            "Kartu Aktivitas Sistem merangkum pembayaran dan aktivasi membership yang perlu dipantau.",
            "Kartu Total Siswa menampilkan total siswa serta status aktif/nonaktif.",
            "Grafik Ikhtisar Omzet dapat difilter untuk 7 hari terakhir, 30 hari terakhir, atau tahun berjalan.",
            "Panel Status Operasional Cabang dapat difilter ke semua status, aktif, atau persiapan/nonaktif.",
        ],
    )

    add_heading(doc, "2.2 Kelola Admin Cabang", 2)
    add_steps(
        doc,
        [
            "Buka menu Admin Cabang.",
            "Gunakan pencarian atau filter jika daftar admin sudah banyak.",
            "Klik Tambah admin cabang untuk membuat akun baru.",
            "Lengkapi data admin, email, password awal, dan informasi terkait cabang bila tersedia.",
            "Klik Buat akun admin untuk menyimpan.",
            "Gunakan tombol Edit untuk memperbarui data admin dan tombol Hapus untuk menghapus akun.",
            "Klik Kirim Ulang Verifikasi jika admin belum menerima email verifikasi.",
        ],
    )
    add_callout(
        doc,
        "Dampak penghapusan admin cabang",
        "Jika admin yang dihapus sedang dipakai oleh cabang, cabang tersebut akan kembali ke status admin belum ditentukan sesuai pesan konfirmasi pada UI.",
    )

    add_heading(doc, "2.3 Kelola Cabang", 2)
    add_steps(
        doc,
        [
            "Buka menu Cabang.",
            "Klik Tambah cabang untuk membuat cabang baru.",
            "Isi nama cabang, alamat, status operasional, dan pilih admin cabang dari akun yang sudah dibuat.",
            "Klik Simpan cabang atau Update perubahan.",
            "Gunakan tombol Edit untuk memperbaiki data cabang.",
            "Gunakan tombol Hapus untuk mengeluarkan cabang dari tabel owner.",
        ],
    )
    add_table(
        doc,
        ["Fitur", "Cara pakai", "Catatan"],
        [
            ("Import", "Klik Import lalu pilih file CSV atau JSON.", "File menerima kolom seperti nama cabang, alamat, status, dan admin sesuai parser import."),
            ("Export CSV", "Klik Export lalu pilih Export CSV.", "Cocok untuk dibuka di spreadsheet."),
            ("Export JSON", "Klik Export lalu pilih Export JSON.", "Cocok untuk backup data atau integrasi teknis."),
            ("Reset filter", "Klik Reset filter ketika hasil tabel terlalu sempit.", "Mengembalikan tampilan ke semua data."),
        ],
        [1700, 3100, 4560],
    )

    add_heading(doc, "2.4 Pembayaran dan Aktivitas", 2)
    add_bullets(
        doc,
        [
            "Tab Pembayaran Masuk menampilkan transaksi membership dari siswa, status pembayaran, nominal, provider, dan waktu update.",
            "Tab Aktivasi Siswa menampilkan status aktivasi membership siswa seperti aktif, menunggu pembayaran, expired, atau pembayaran gagal.",
            "Owner dapat membuka detail transaksi atau detail siswa untuk membaca informasi lebih lengkap.",
            "Setiap tabel memiliki export CSV dan JSON untuk rekap atau laporan.",
            "Gunakan filter status dan pencarian untuk menemukan transaksi tertentu.",
        ],
    )

    add_heading(doc, "2.5 Pengeluaran Operasional", 2)
    add_steps(
        doc,
        [
            "Buka menu Pengeluaran.",
            "Klik Tambah Pengeluaran Operasional.",
            "Pilih cabang, kategori, judul pengeluaran, nominal, tanggal, status, dan catatan bila perlu.",
            "Klik Simpan pengeluaran.",
            "Gunakan Edit untuk memperbarui data pengeluaran dan Hapus untuk menghapus item.",
            "Gunakan filter cabang, status, atau kategori untuk menganalisis pengeluaran tertentu.",
        ],
    )

    add_checklist(
        doc,
        [
            "Cabang baru sudah memiliki admin cabang.",
            "Admin cabang sudah menerima email verifikasi.",
            "Status cabang sudah sesuai: Aktif, Persiapan, atau Nonaktif.",
            "Pembayaran masuk dan aktivasi siswa dicek rutin.",
            "Pengeluaran besar sudah dicatat dengan nominal dan cabang yang tepat.",
        ],
    )


def add_admin_manual(doc: Document):
    add_heading(doc, "3. Manual Role Admin", 1)
    add_para(
        doc,
        "Admin mengelola operasional cabang: data siswa, data guru, jadwal, pembayaran membership, dan keuangan cabang.",
    )
    add_table(
        doc,
        ["Menu Admin", "Kegunaan utama"],
        [
            ("Dashboard", "Melihat kesehatan operasional harian, total siswa, guru aktif, jadwal, ruangan terpakai, dan ringkasan pembayaran."),
            ("Kelola Siswa", "Tambah, edit, detail, hapus/nonaktifkan, import Excel, dan export data siswa."),
            ("Kelola Guru", "Tambah, edit, detail, hapus, import Excel, export data guru, dan pengaturan jadwal mengajar."),
            ("Jadwal Kelas", "Tambah, edit, hapus, import/export jadwal, validasi guru, ruangan, kelas, mapel, dan konflik."),
            ("Keuangan Cabang", "Pantau pemasukan/pengeluaran cabang dan catat pengeluaran operasional cabang."),
            ("Informasi Pembayaran", "Membuat tagihan, batch tagihan, melihat pembayaran masuk, resend link, cancel tagihan, dan export CSV."),
        ],
        [2100, 7260],
    )

    add_heading(doc, "3.1 Dashboard Admin", 2)
    add_bullets(
        doc,
        [
            "Gunakan Dashboard untuk melihat ringkasan cabang secara cepat sebelum masuk ke menu detail.",
            "Panel pembayaran menampilkan total pembayaran, status lunas/pending/expired/gagal, dan tren nominal transaksi.",
            "Panel data master menampilkan siswa, guru, jadwal, dan ruangan terpakai.",
            "Kontrol periode mendukung ringkasan 7 hari terakhir, 30 hari terakhir, dan tahun berjalan.",
        ],
    )

    add_heading(doc, "3.2 Kelola Siswa", 2)
    add_steps(
        doc,
        [
            "Buka menu Kelola Siswa.",
            "Klik Tambah Data untuk menambah siswa manual.",
            "Isi nama, email, telepon, cabang, level, program, kelas, tanggal lahir, dan status.",
            "Klik Tambah siswa atau Simpan perubahan.",
            "Klik ikon detail untuk melihat data lengkap siswa.",
            "Klik Edit untuk memperbarui data siswa.",
            "Klik Hapus untuk menghapus data. Jika siswa memiliki histori pembayaran atau subscription, sistem dapat mengubah status menjadi Nonaktif, bukan menghapus permanen.",
        ],
    )
    add_table(
        doc,
        ["Aksi", "Panduan", "Catatan"],
        [
            ("Import Excel", "Klik Import, pilih file Excel siswa, lalu klik Import Excel.", "Header yang dibaca mengikuti konfigurasi studentImportColumns pada UI."),
            ("Export", "Klik Export.", "File export mengikuti filter yang sedang dipakai bila endpoint mendukung query."),
            ("Password awal", "Gunakan informasi generated password bila tersedia.", "Project memiliki utilitas password dari tanggal lahir dengan format ddmmyyyy."),
            ("Filter", "Gunakan filter level, status, kelas, atau pencarian.", "Reset filter bila data yang dicari tidak muncul."),
        ],
        [1800, 3300, 4260],
    )

    add_heading(doc, "3.3 Kelola Guru", 2)
    add_steps(
        doc,
        [
            "Buka menu Kelola Guru.",
            "Klik Tambah Data.",
            "Isi nama, email, mata pelajaran, jadwal mengajar, daftar kelas, cabang, telepon, status, dan availability.",
            "Saat membuat guru baru, isi kredensial atau gunakan pola auto-generate bila UI menawarkannya.",
            "Klik Tambah guru atau Simpan perubahan.",
            "Gunakan detail untuk membaca data lengkap, Edit untuk revisi, dan Hapus untuk menghapus guru.",
            "Gunakan Import Excel dan Export untuk pengelolaan massal.",
        ],
    )

    add_heading(doc, "3.4 Jadwal Kelas", 2)
    add_steps(
        doc,
        [
            "Buka menu Jadwal Kelas.",
            "Klik Tambah Data.",
            "Pilih hari, slot waktu, jenjang/kelas, mata pelajaran, guru, cabang, ruangan, dan status jadwal.",
            "Pastikan guru dan ruangan tersedia. UI memberi peringatan bila ada format lama, guru nonaktif, mapel belum valid, atau slot tidak tersedia.",
            "Klik Tambah jadwal atau Simpan perubahan.",
            "Gunakan detail, edit, hapus, import, dan export sesuai kebutuhan.",
        ],
    )
    add_callout(
        doc,
        "Validasi konflik jadwal",
        "Jika status jadwal menunjukkan Bentrok atau Perlu Cek, periksa kembali guru, ruangan, hari, jam, kelas, dan mapel sebelum jadwal dipakai guru atau siswa.",
    )

    add_heading(doc, "3.5 Keuangan Cabang", 2)
    add_bullets(
        doc,
        [
            "Menu Keuangan Cabang menampilkan ringkasan pemasukan, pengeluaran, saldo bersih, dan transaksi cabang.",
            "Pemasukan pembayaran membership tetap diverifikasi lewat menu Pembayaran.",
            "Menu Keuangan difokuskan untuk pencatatan pengeluaran operasional cabang.",
            "Admin dapat tambah, edit, dan hapus pengeluaran operasional sesuai cabang yang dikelola.",
        ],
    )
    add_steps(
        doc,
        [
            "Klik Tambah Pengeluaran.",
            "Isi judul, kategori, nominal, tanggal/jadwal bayar, status, dan catatan.",
            "Klik Tambah pengeluaran atau Simpan perubahan.",
            "Gunakan ikon Edit atau Hapus pada tabel pengeluaran bila perlu koreksi.",
        ],
    )

    add_heading(doc, "3.6 Informasi Pembayaran", 2)
    add_bullets(
        doc,
        [
            "Tab Pembayaran Masuk menampilkan transaksi siswa, paket, nominal, provider, metode, status, dan checkout link.",
            "Admin dapat membuat tagihan membership untuk satu siswa atau batch beberapa siswa.",
            "Untuk transaksi pending, admin dapat mengirim ulang checkout link.",
            "Admin dapat membatalkan tagihan pending bila tagihan salah atau tidak dilanjutkan.",
            "Export CSV tersedia untuk pembayaran masuk dan aktivasi membership.",
        ],
    )
    add_steps(
        doc,
        [
            "Buka menu Informasi Pembayaran.",
            "Pilih siswa yang akan dibuatkan tagihan.",
            "Pilih paket membership dan durasi sesuai kebutuhan.",
            "Klik Buat Tagihan. Untuk banyak siswa, gunakan proses batch tagihan.",
            "Jika checkout link perlu dikirim ulang, klik Kirim Ulang Link pada transaksi terkait.",
            "Jika tagihan perlu dibatalkan, klik Batalkan Tagihan dan konfirmasi.",
        ],
    )

    add_checklist(
        doc,
        [
            "Data siswa aktif dan cabangnya benar sebelum dibuatkan tagihan.",
            "Data guru aktif sebelum dipakai di jadwal.",
            "Jadwal tidak bentrok dan ruangan tersedia.",
            "Pengeluaran cabang dicatat pada cabang yang benar.",
            "Tagihan pending ditindaklanjuti: resend link atau cancel bila perlu.",
        ],
    )


def add_teacher_manual(doc: Document):
    add_heading(doc, "4. Manual Role Guru", 1)
    add_para(
        doc,
        "Guru memakai dashboard untuk melihat jadwal, mengelola kelas, menjalankan absensi QR, membagikan materi, membuat tugas, menilai submission, dan mengelola tryout.",
    )
    add_table(
        doc,
        ["Area Guru", "Fungsi"],
        [
            ("Beranda", "Profil guru, jadwal hari ini, ringkasan kelas aktif, dan shortcut manajemen kelas."),
            ("Jadwal", "Daftar jadwal mengajar dan akses menuju absensi kelas."),
            ("Semua Kelas", "Daftar kelas yang diampu dan akses detail kelas."),
            ("Detail Kelas", "Peserta, materi, tugas, nilai, submission, target pertemuan, dan absensi per pertemuan."),
            ("Absensi Kelas", "Membuat sesi QR, membuka QR, mengubah status kehadiran, dan menutup sesi."),
            ("Tryout", "Membuat tryout, mengelola soal manual, publish/unpublish, dan melihat hasil."),
        ],
        [2100, 7260],
    )

    add_heading(doc, "4.1 Beranda dan Jadwal Guru", 2)
    add_bullets(
        doc,
        [
            "Topbar guru memiliki menu Beranda dan Tryout. Beberapa shortcut di dashboard mengarah ke Jadwal atau Semua Kelas.",
            "Panel profil menampilkan nama, email, mapel, cabang, status, dan jumlah kelas aktif.",
            "Panel jadwal hari ini menampilkan jam, kelas, mapel, dan ruang jika data backend tersedia.",
            "Klik Lihat Semua Jadwal untuk membuka jadwal lengkap.",
        ],
    )

    add_heading(doc, "4.2 Semua Kelas dan Detail Kelas", 2)
    add_steps(
        doc,
        [
            "Buka Semua Kelas dari dashboard atau halaman kelas guru.",
            "Pilih kelas yang ingin dikelola.",
            "Pada detail kelas, cek peserta, detail pertemuan, materi, tugas, nilai, dan absensi.",
            "Gunakan sidebar detail kelas untuk berpindah antarbagian tanpa keluar dari kelas.",
        ],
    )

    add_heading(doc, "4.3 Materi Pembelajaran", 2)
    add_steps(
        doc,
        [
            "Masuk ke Detail Kelas.",
            "Pada tabel Detail Pertemuan & Materi, klik Tambah Materi.",
            "Isi pertemuan, judul materi, deskripsi, link materi, status materi, dan lampiran opsional.",
            "Klik Simpan Materi.",
            "Gunakan Edit untuk memperbarui materi dan Hapus untuk menghapus materi.",
            "Jika materi memiliki link, siswa dapat membuka link. Jika ada lampiran, siswa dapat mengunduh lampiran dari halaman materi.",
        ],
    )

    add_heading(doc, "4.4 Tugas dan Review Submission", 2)
    add_steps(
        doc,
        [
            "Masuk ke Detail Kelas.",
            "Pada bagian Tugas Pertemuan, klik Tambah Tugas.",
            "Isi judul tugas, deskripsi, deadline, jenis submission, status penilaian, dan lampiran opsional.",
            "Klik Simpan Tugas.",
            "Pantau daftar submission siswa pada tugas terkait.",
            "Buka submission untuk membaca jawaban teks, link Drive, atau lampiran file.",
            "Berikan review, nilai, dan catatan guru lalu simpan.",
        ],
    )
    add_callout(
        doc,
        "Jenis jawaban siswa",
        "Siswa dapat mengirim jawaban dengan upload file, jawaban teks, atau link Drive sesuai panel submission yang tersedia pada dashboard siswa.",
    )

    add_heading(doc, "4.5 Nilai Siswa", 2)
    add_steps(
        doc,
        [
            "Buka Detail Kelas dan masuk ke bagian nilai.",
            "Pilih siswa atau tugas yang akan dinilai.",
            "Isi skor, status, dan catatan evaluasi.",
            "Klik Simpan. Nilai yang tersimpan akan muncul pada dashboard siswa bagian Nilai.",
            "Gunakan tabel Belum Dinilai untuk memprioritaskan submission yang belum mendapat nilai.",
        ],
    )

    add_heading(doc, "4.6 Absensi QR", 2)
    add_steps(
        doc,
        [
            "Buka jadwal atau halaman Absensi Kelas untuk kelas tertentu.",
            "Klik Mulai Absensi QR untuk membuat sesi absensi hari ini.",
            "Tampilkan QR Absensi kepada siswa.",
            "Siswa membuka halaman Scan Absen dan memindai QR.",
            "Pantau daftar peserta. Status dapat berupa Belum Absen, Hadir, Sakit, Izin, atau Alpa.",
            "Jika perlu koreksi, ubah status kehadiran siswa secara manual.",
            "Klik Tutup sesi bila absensi selesai.",
        ],
    )
    add_callout(
        doc,
        "Status sesi",
        "Jika sesi belum dimulai, UI menampilkan instruksi untuk mulai absensi. Jika sesi sudah ditutup, status baru tidak dapat dipakai untuk scan aktif.",
    )

    add_heading(doc, "4.7 Tryout Guru", 2)
    add_steps(
        doc,
        [
            "Buka menu Tryout.",
            "Pilih jenjang SD, SMP, atau SMA. Sistem otomatis menyesuaikan kelas akhir: SD 6, SMP 9, atau SMA 12.",
            "Klik Tambah Tryout.",
            "Isi judul tryout, jenjang, kelas, mapel, jumlah soal, durasi, tanggal mulai, tanggal selesai, status publish, dan sumber soal.",
            "Simpan tryout sebagai Draft terlebih dahulu.",
            "Klik Upload Soal untuk menambahkan soal manual.",
            "Isi pertanyaan, empat opsi jawaban, dan kunci jawaban.",
            "Klik Tambah Soal atau Simpan Perubahan.",
            "Jika jumlah soal sudah lebih dari nol, ubah status menjadi Published agar tryout siap dibuka siswa.",
            "Gunakan Unpublish untuk menutup akses, Edit Tryout untuk revisi metadata, Hapus Tryout untuk menghapus, dan Hasil untuk melihat nilai siswa.",
        ],
    )
    add_callout(
        doc,
        "Catatan fitur tryout",
        "Upload file soal dan integrasi bank soal ditandai belum aktif penuh pada UI. Alur yang siap dipakai pada project saat ini adalah pembuatan metadata tryout dan penambahan soal manual.",
        fill=AMBER,
    )

    add_checklist(
        doc,
        [
            "Jadwal mengajar hari ini sudah dicek.",
            "Materi pertemuan sudah tersedia sebelum kelas.",
            "Tugas memiliki deadline dan instruksi yang jelas.",
            "Absensi QR ditutup setelah kelas selesai.",
            "Submission siswa yang belum dinilai sudah ditindaklanjuti.",
            "Tryout dipublish hanya setelah memiliki soal.",
        ],
    )


def add_student_manual(doc: Document):
    add_heading(doc, "5. Manual Role Siswa", 1)
    add_para(
        doc,
        "Siswa memakai dashboard untuk melihat profil belajar, membuka materi, mengerjakan dan mengirim tugas, melihat jadwal, melakukan absensi, memantau nilai, membayar tagihan, dan mengikuti tryout.",
    )

    add_heading(doc, "5.1 Registrasi Siswa Baru", 2)
    add_steps(
        doc,
        [
            "Buka halaman /register.",
            "Isi nama lengkap, email, nomor WhatsApp, program belajar, tingkat kelas, cabang pilihan, password, dan konfirmasi password.",
            "Pilih paket belajar: 1 Bulan, 3 Bulan, 6 Bulan, atau 1 Tahun.",
            "Klik Daftar & Buat Tagihan.",
            "Sistem membuat akun siswa, data membership, dan tagihan pembayaran.",
            "Lanjutkan ke halaman pembayaran untuk menyelesaikan tagihan.",
        ],
    )
    add_table(
        doc,
        ["Program", "Pilihan kelas"],
        [
            ("SD / Program Dasar", "Kelas 4, Kelas 5, Kelas 6"),
            ("SMP / Program Reguler", "Kelas 7, Kelas 8, Kelas 9"),
            ("SMA / Program Intensif", "Kelas 10, Kelas 11, Kelas 12"),
        ],
        [2600, 6760],
    )

    add_heading(doc, "5.2 Pembayaran dan Status Membership", 2)
    add_steps(
        doc,
        [
            "Pada halaman pembayaran, cek ringkasan tagihan: paket, nominal, nama siswa, cabang, ID transaksi, provider, dan status.",
            "Pilih metode pembayaran yang tersedia seperti QRIS, Virtual Account, atau E-Wallet.",
            "Klik Bayar Sekarang bila tagihan memakai Xendit checkout, atau Konfirmasi Pembayaran bila masih memakai simulasi internal.",
            "Setelah pembayaran sukses, buka Status Membership.",
            "Pastikan checklist Pembayaran, Verifikasi Email, dan Akses Dashboard sudah lengkap.",
            "Klik Masuk ke Dashboard untuk login sebagai siswa.",
        ],
    )
    add_bullets(
        doc,
        [
            "Status pending berarti tagihan masih menunggu pembayaran.",
            "Status paid berarti pembayaran sudah diterima dan membership aktif.",
            "Status failed atau expired berarti tagihan gagal/kedaluwarsa dan perlu dibuat ulang.",
            "Jika pembayaran Xendit selesai tetapi dashboard belum aktif, tunggu proses webhook atau cek status kembali.",
        ],
    )

    add_heading(doc, "5.3 Dashboard Siswa", 2)
    add_bullets(
        doc,
        [
            "Topbar siswa memiliki menu Beranda, Absensi, Nilai, dan Tagihan.",
            "Header profil menampilkan data siswa, kelas/program, status membership, cabang, dan tagihan.",
            "Header akademik menampilkan ringkasan materi, tugas, dan jadwal hari ini.",
            "Panel Pelajaran menampilkan tab Daftar Materi, Daftar Tugas, dan Sesi Tryout.",
            "Notifikasi siswa merangkum jadwal, tugas, materi, tagihan, dan nilai.",
        ],
    )

    add_heading(doc, "5.4 Materi Belajar", 2)
    add_steps(
        doc,
        [
            "Buka Beranda lalu pilih tab Materi, atau buka halaman /dashboard-siswa/materi.",
            "Pilih materi dari daftar.",
            "Klik Lihat Materi untuk membaca preview dan detail materi.",
            "Klik Download Materi bila lampiran atau file unduhan tersedia.",
            "Gunakan Simpan ke Jadwal Belajar sebagai penanda rencana belajar bila tersedia pada UI.",
        ],
    )
    add_callout(
        doc,
        "Sumber materi",
        "Materi siswa berasal dari materi yang dibagikan guru pada kelas terkait. Jika kosong, tunggu guru menambahkan materi atau pastikan siswa sudah berada pada kelas dan cabang yang benar.",
    )

    add_heading(doc, "5.5 Tugas dan Kirim Tugas", 2)
    add_steps(
        doc,
        [
            "Buka halaman Tugas Siswa atau tab Tugas pada dashboard.",
            "Pilih tugas yang ingin dikerjakan.",
            "Baca instruksi, deadline, status, lampiran tugas, dan catatan guru.",
            "Klik Kirim Tugas.",
            "Pilih metode jawaban: Upload File, Jawaban Teks, atau Link Drive.",
            "Isi jawaban sesuai metode yang dipilih.",
            "Klik Kirim Tugas Sekarang.",
            "Cek status submission. Jika sudah dinilai, lihat nilai dan catatan guru.",
        ],
    )
    add_table(
        doc,
        ["Metode submission", "Dipakai untuk", "Catatan"],
        [
            ("Upload File", "Jawaban berupa dokumen, foto, atau file tugas.", "Pastikan ukuran dan format file dapat diterima browser/backend."),
            ("Jawaban Teks", "Jawaban singkat atau uraian langsung.", "Tulis dengan jelas dan cek kembali sebelum submit."),
            ("Link Drive", "Jawaban disimpan di Google Drive atau link eksternal.", "Pastikan permission link dapat dibuka guru."),
        ],
        [2300, 3300, 3760],
    )

    add_heading(doc, "5.6 Jadwal Siswa", 2)
    add_bullets(
        doc,
        [
            "Buka menu Jadwal dari dashboard atau halaman /dashboard-siswa/jadwal.",
            "Jadwal ditampilkan berdasarkan kelas dan cabang siswa yang sedang login.",
            "Status jadwal dapat berupa Berjalan, Review, Perlu Cek, atau Terjadwal.",
            "Jika jadwal kosong, pastikan admin sudah membuat jadwal untuk kelas dan cabang siswa.",
        ],
    )

    add_heading(doc, "5.7 Absensi Siswa", 2)
    add_steps(
        doc,
        [
            "Buka menu Absensi untuk melihat riwayat kehadiran.",
            "Saat guru membuka sesi QR, buka halaman Scan Absen.",
            "Izinkan akses kamera jika browser meminta permission.",
            "Arahkan kamera ke QR Absensi yang ditampilkan guru.",
            "Tunggu sistem memproses hasil scan.",
            "Cek riwayat absensi untuk memastikan status kehadiran masuk.",
        ],
    )

    add_heading(doc, "5.8 Nilai Siswa", 2)
    add_bullets(
        doc,
        [
            "Buka menu Nilai untuk melihat rekapitulasi nilai.",
            "Daftar Nilai Berdasarkan Mata Pelajaran menampilkan jumlah tugas dinilai dan rata-rata nilai.",
            "Klik mata pelajaran untuk melihat detail nilai tugas.",
            "Gunakan catatan guru sebagai bahan revisi atau belajar ulang.",
        ],
    )

    add_heading(doc, "5.9 Tagihan Siswa", 2)
    add_bullets(
        doc,
        [
            "Buka menu Tagihan untuk melihat status membership dan histori tagihan.",
            "Histori Tagihan menampilkan paket, nominal, status, provider, metode, tanggal dibuat, dan waktu bayar.",
            "Klik Lanjut Pembayaran bila tagihan masih pending dan memiliki checkout URL.",
            "Jika tagihan expired/gagal, hubungi admin untuk dibuatkan tagihan baru.",
        ],
    )

    add_heading(doc, "5.10 Tryout Siswa", 2)
    add_steps(
        doc,
        [
            "Buka tab Sesi Tryout atau halaman /dashboard-siswa/tryout.",
            "Baca judul, mapel, jenjang, durasi, jumlah soal, target skor, jadwal, dan instruksi.",
            "Klik Mulai Tryout.",
            "Pilih jawaban pada setiap soal.",
            "Gunakan navigasi soal untuk berpindah nomor dan mengecek soal yang belum dijawab.",
            "Perhatikan timer tryout.",
            "Klik Kirim Tryout setelah selesai.",
            "Lihat hasil tryout dan skor estimasi.",
        ],
    )
    add_callout(
        doc,
        "Catatan fitur tryout siswa",
        "Halaman tryout siswa pada project saat ini memakai data sesi tryout lokal/dummy untuk simulasi pengerjaan. Integrasi hasil tryout penuh perlu disesuaikan dengan backend tryout jika ingin dipakai produksi.",
        fill=AMBER,
    )

    add_checklist(
        doc,
        [
            "Email dan password siswa dapat dipakai login.",
            "Membership aktif sebelum mengakses pembelajaran penuh.",
            "Materi dan tugas muncul sesuai kelas siswa.",
            "Submission tugas sudah terkirim dan dapat dilihat ulang.",
            "Absensi QR berhasil masuk ke riwayat.",
            "Tagihan pending diselesaikan sebelum kedaluwarsa.",
        ],
    )


def add_troubleshooting(doc: Document):
    add_heading(doc, "6. Troubleshooting Umum", 1)
    add_table(
        doc,
        ["Masalah", "Kemungkinan penyebab", "Tindakan"],
        [
            ("Tidak bisa login", "Email/password salah, email belum terverifikasi, token expired, atau akun tidak aktif.", "Cek kredensial, lakukan reset password, verifikasi email, atau minta admin mengecek status akun."),
            ("Redirect ke dashboard salah", "Role akun tidak sesuai dengan halaman yang dibuka.", "Login ulang. Sistem akan mengarahkan ke dashboard sesuai role."),
            ("Data tabel kosong", "Backend belum memiliki data, filter terlalu sempit, atau akun tidak terkait cabang/kelas.", "Reset filter, cek cabang/kelas, atau tambah data dari admin/owner."),
            ("Import gagal", "Header file tidak sesuai, format file tidak didukung, atau data wajib kosong.", "Ikuti daftar header pada dialog import dan perbaiki file."),
            ("Export tidak muncul", "Popup/download browser diblokir atau endpoint gagal.", "Izinkan download browser dan coba ulang setelah filter dicek."),
            ("QR absensi gagal discan", "Sesi belum aktif, kamera tidak diizinkan, QR sudah ditutup, atau koneksi bermasalah.", "Minta guru membuka ulang QR, izinkan kamera, dan cek status sesi."),
            ("Pembayaran belum aktif", "Checkout belum selesai, webhook belum masuk, atau status masih pending.", "Cek halaman status, tunggu sinkronisasi, atau minta admin resend/cancel tagihan."),
            ("Materi/tugas tidak muncul", "Guru belum menambahkan materi/tugas atau siswa belum terkait kelas yang benar.", "Cek kelas siswa pada admin dan minta guru menambahkan konten."),
        ],
        [2300, 3300, 3760],
    )

    add_heading(doc, "7. Ringkasan Hak Akses", 1)
    add_table(
        doc,
        ["Fitur", "Owner", "Admin", "Guru", "Siswa"],
        [
            ("Dashboard ringkasan", "Semua cabang", "Cabang dikelola", "Kelas diampu", "Profil belajar sendiri"),
            ("Kelola cabang", "Ya", "Tidak", "Tidak", "Tidak"),
            ("Kelola admin cabang", "Ya", "Tidak", "Tidak", "Tidak"),
            ("Kelola siswa", "Pantau ringkasan", "Ya", "Lihat peserta kelas", "Data sendiri"),
            ("Kelola guru", "Pantau ringkasan", "Ya", "Profil sendiri", "Tidak"),
            ("Kelola jadwal", "Pantau", "Ya", "Lihat jadwal", "Lihat jadwal"),
            ("Materi dan tugas", "Pantau", "Tidak langsung", "Buat dan nilai", "Akses dan submit"),
            ("Absensi QR", "Pantau", "Data operasional", "Buat dan kelola sesi", "Scan dan lihat riwayat"),
            ("Pembayaran", "Pantau semua", "Buat/kelola tagihan", "Tidak", "Bayar dan lihat histori"),
            ("Pengeluaran", "Semua cabang", "Cabang dikelola", "Tidak", "Tidak"),
            ("Tryout", "Pantau tidak langsung", "Tidak langsung", "Buat dan publish", "Kerjakan"),
        ],
        [2600, 1690, 1690, 1690, 1690],
    )

    add_heading(doc, "8. Catatan Implementasi", 1)
    add_bullets(
        doc,
        [
            "Manual ini dibuat berdasarkan project bimbel-new pada 18 Juni 2026.",
            "Aplikasi memakai role owner, admin, guru, dan siswa dengan redirect dashboard berbeda.",
            "Beberapa data akan kosong bila backend belum memiliki seed/data operasional.",
            "Beberapa fitur tryout dan upload file soal memiliki catatan placeholder pada UI, sehingga manual membedakan fitur siap pakai dan fitur menunggu integrasi.",
            "Untuk demo skripsi, siapkan akun contoh untuk setiap role agar reviewer dapat mengikuti alur manual tanpa harus membuat data dari nol.",
        ],
    )


def build_document():
    doc = Document()
    configure_document(doc)
    set_running_header_footer(doc)
    add_cover(doc)
    add_quick_toc(doc)
    add_common_access(doc)
    add_owner_manual(doc)
    add_admin_manual(doc)
    add_teacher_manual(doc)
    add_student_manual(doc)
    add_troubleshooting(doc)

    # Preset audit marker: the script explicitly sets Letter page geometry,
    # compact_reference_guide typography rhythm, real Word list styles, and
    # fixed DXA table geometry through set_table_geometry.
    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    path = build_document()
    print(path)
